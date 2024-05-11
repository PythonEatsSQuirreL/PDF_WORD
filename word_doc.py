from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('wp6782679.png', width=Inches(1.25))

records = (
    (3, '101', 'Carrots'),
    (7, '422', 'Eggs'),
    (4, '631', 'Ham, Butter, Oranges, and Plumbs')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')

###########################docx_to_pdf######################
from docx2pdf import convert

# Converting docx present in the same folder as the python file
convert("demo.docx")
 
# Converting docx specifying both the input and output paths
#convert("Folder\GFG_1.docx", "Other_Folder\Mine.pdf")
 

#Convert an entire folder full of docx files to pdf files
#convert("Folder\\")