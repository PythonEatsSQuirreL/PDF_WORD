
import pdfplumber as ppl
import os
from docx import Document

doc = Document()

pdf_path = r'C:\Users\PsychO\Desktop\yt\demo.pdf'
gen_path = os.path.sep.join(pdf_path.split(os.path.sep)[:-1])

pdf_file = pdf_path.split(os.path.sep)[-1]
output_path = pdf_file.split('.')[0] + ' backup.docx'

p = doc.add_paragraph()


with ppl.open(pdf_path) as pdf:
    for page in pdf.pages:
        textt = page.extract_text()
        #the data from the entire pdf
        print(textt)
        p.add_run(textt)



doc.save(gen_path + "\\" + output_path)


